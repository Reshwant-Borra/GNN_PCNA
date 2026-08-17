"""Assemble the manuscript into a competition-format Word document.

Single-column, serif, numbered sections, real figures embedded with numbered
captions, and a references section. A visible draft notice keeps the
human-in-the-loop framing (the document is an auto-generated draft that must be
reviewed before submission).
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from paper_engine.manuscript.bibliography import Reference
from paper_engine.manuscript.section_writer import SectionResult


def _set_base_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Cambria"
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)


def _paragraphs(text: str) -> List[str]:
    chunks = re.split(r"\n\s*\n", text.strip())
    if len(chunks) == 1:
        chunks = [c for c in text.split("\n") if c.strip()]
    return [c.strip() for c in chunks if c.strip()]


def _add_figure(doc: Document, fig: dict, number: int) -> None:
    path = Path(fig["path"])
    if not path.exists():
        return
    # Square figures (e.g. DCCM) get a smaller width.
    width = Inches(4.2) if "dccm" in fig.get("figure_id", "") else Inches(5.9)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(f"Figure {number}. {fig.get('caption', '')}")
    cr.font.size = Pt(9.5)
    cr.italic = True


def assign_figure_numbers(sections: List[SectionResult],
                          figure_map: Dict[str, dict]) -> Dict[str, int]:
    """Number figures by first appearance across the section order."""
    numbers: Dict[str, int] = {}
    n = 0
    for sec in sections:
        for fid in sec.figures:
            if fid in figure_map and Path(figure_map[fid]["path"]).exists() and fid not in numbers:
                n += 1
                numbers[fid] = n
    return numbers


def build_docx(meta: dict, sections: List[SectionResult], figure_map: Dict[str, dict],
               references: List[Reference], out_path: Path) -> Path:
    doc = Document()
    _set_base_style(doc)

    # --- Title block ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(meta["title"])
    tr.bold = True
    tr.font.size = Pt(18)
    if meta.get("subtitle"):
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sub.add_run(meta["subtitle"])
        sr.italic = True
        sr.font.size = Pt(12)
    byline = doc.add_paragraph()
    byline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = byline.add_run(f"{meta.get('author', '[Author]')}    ·    {meta.get('date', '')}")
    br.font.size = Pt(11)

    # --- Draft notice (integrity) ---
    notice = doc.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = notice.add_run(meta.get("draft_notice", ""))
    nr.italic = True
    nr.font.size = Pt(8.5)
    nr.font.color.rgb = RGBColor(0x88, 0x44, 0x00)

    figure_numbers = assign_figure_numbers(sections, figure_map)

    # --- Sections ---
    for sec in sections:
        level = 0 if sec.section_id == "abstract" else 1
        doc.add_heading(sec.heading, level=1 if level else 1)
        for para in _paragraphs(sec.text):
            doc.add_paragraph(para)
        # Embed this section's figures right after its prose.
        for fid in sec.figures:
            if fid in figure_numbers:
                _add_figure(doc, {**figure_map[fid], "figure_id": fid}, figure_numbers[fid])

    # --- References ---
    if references:
        doc.add_heading("References", level=1)
        for r in references:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.first_line_indent = Inches(-0.3)
            run = p.add_run(f"[{r.idx}] {r.formatted()}")
            run.font.size = Pt(9.5)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
