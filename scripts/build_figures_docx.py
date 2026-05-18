"""
Build a figures-only Word document for journal assembly.
Generates architecture diagrams + assembles all meaningful data figures.
Output: docs/GNN_PCNA_Figures.docx
"""
import io
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
import numpy as np

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT    = Path(__file__).parent.parent
DOCS    = ROOT / "docs"
FIGS    = ROOT / "results" / "figures"
RESULTS = ROOT / "results"
DATA    = ROOT / "data" / "results"
OUT     = DOCS / "GNN_PCNA_Figures_v2.docx"

NAVY  = "#1F497D"
BLUE  = "#2E74B5"
LGRAY = "#F2F2F2"
DGRAY = "#595959"

# ──────────────────────────────────────────────────────────────────────────────
# Architecture diagram helpers
# ──────────────────────────────────────────────────────────────────────────────

def box(ax, cx, cy, w, h, label, sublabel=None,
        fc="#DDEEFF", ec=BLUE, fontsize=9, bold=False):
    rect = FancyBboxPatch(
        (cx - w/2, cy - h/2), w, h,
        boxstyle="round,pad=0.02", linewidth=1.2,
        edgecolor=ec, facecolor=fc, zorder=3
    )
    ax.add_patch(rect)
    weight = "bold" if bold else "normal"
    if sublabel:
        ax.text(cx, cy + h*0.12, label, ha="center", va="center",
                fontsize=fontsize, fontweight=weight, color="#1a1a1a", zorder=4)
        ax.text(cx, cy - h*0.18, sublabel, ha="center", va="center",
                fontsize=fontsize - 1.5, color=DGRAY, zorder=4, style="italic")
    else:
        ax.text(cx, cy, label, ha="center", va="center",
                fontsize=fontsize, fontweight=weight, color="#1a1a1a", zorder=4)

def arrow(ax, x1, y1, x2, y2, col="#444444"):
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle="-|>", color=col,
                                lw=1.3, mutation_scale=12), zorder=2)

def bracket_label(ax, x, y1, y2, label, col=BLUE):
    ax.annotate("", xy=(x, y2), xytext=(x, y1),
                arrowprops=dict(arrowstyle="-", color=col, lw=1.5,
                                connectionstyle="bar,fraction=0.0"), zorder=2)
    ax.text(x + 0.04, (y1 + y2) / 2, label, va="center", ha="left",
            fontsize=8, color=col, fontweight="bold")

# ──────────────────────────────────────────────────────────────────────────────
# Figure 1 — PocketGNNXL Model Architecture
# ──────────────────────────────────────────────────────────────────────────────

def make_arch_diagram() -> bytes:
    fig, ax = plt.subplots(figsize=(13, 8.5))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    # Title
    ax.text(0.5, 0.97, "PocketGNNXL — Model Architecture",
            ha="center", va="top", fontsize=13, fontweight="bold", color=NAVY)

    # ── Input nodes ──────────────────────────────────────────────────────────
    bw, bh = 0.18, 0.06

    # Two input streams
    box(ax, 0.18, 0.86, bw, bh,
        "Hand-crafted features", "40-dim per residue",
        fc="#E8F4FD", ec=BLUE)
    box(ax, 0.50, 0.86, bw, bh,
        "ESM2 embeddings", "480-dim per residue",
        fc="#FFF3E0", ec="#E67E22")

    # Concatenate
    box(ax, 0.34, 0.75, 0.20, bh,
        "Concatenate", "→ 520-dim node features",
        fc="#EEF", ec=NAVY, bold=True)
    arrow(ax, 0.18, 0.83, 0.30, 0.78)
    arrow(ax, 0.50, 0.83, 0.38, 0.78)

    # Edge features (side note)
    box(ax, 0.78, 0.75, 0.16, bh,
        "Edge features", "6-dim (spatial + seq.)",
        fc="#F0FFF0", ec="#27AE60")

    # Pre-encoder
    box(ax, 0.34, 0.64, 0.22, bh,
        "Pre-Encoder", "Linear(520→256→512→768) + LayerNorm",
        fc="#DDE8FF", ec=NAVY, bold=True)
    arrow(ax, 0.34, 0.72, 0.34, 0.67)

    # Branch split
    ax.plot([0.34, 0.34], [0.61, 0.585], color="#444", lw=1.3, zorder=2)
    ax.plot([0.20, 0.50], [0.585, 0.585], color="#444", lw=1.3, zorder=2)
    ax.annotate("", xy=(0.20, 0.565), xytext=(0.20, 0.585),
                arrowprops=dict(arrowstyle="-|>", color="#444", lw=1.3,
                                mutation_scale=12), zorder=2)
    ax.annotate("", xy=(0.50, 0.565), xytext=(0.50, 0.585),
                arrowprops=dict(arrowstyle="-|>", color="#444", lw=1.3,
                                mutation_scale=12), zorder=2)

    # Branch 1 — Spatial
    for i, y in enumerate([0.54, 0.47, 0.40, 0.33, 0.26]):
        box(ax, 0.20, y, 0.23, bh,
            f"GATv2Conv #{i+1}", "hidden=768, heads=8 + BN + Dropout",
            fc="#E3EAF8", ec=BLUE)
        if i > 0:
            arrow(ax, 0.20, y + bh/2 + 0.01, 0.20, y + bh/2 + 0.055)
    # edge input arrow to branch 1
    arrow(ax, 0.70, 0.75, 0.33, 0.545)
    ax.text(0.19, 0.57, "Spatial Branch\n(5× GATv2Conv)", ha="center",
            va="bottom", fontsize=8, color=BLUE, fontweight="bold")

    # Branch 2 — Sequential
    for i, y in enumerate([0.54, 0.47, 0.40, 0.33]):
        box(ax, 0.50, y, 0.23, bh,
            f"GATv2Conv #{i+1}", "hidden=768, heads=8 + BN + Dropout",
            fc="#FDE8D8", ec="#E67E22")
        if i > 0:
            arrow(ax, 0.50, y + bh/2 + 0.01, 0.50, y + bh/2 + 0.055)
    # edge input arrow to branch 2
    arrow(ax, 0.70, 0.75, 0.63, 0.545)
    ax.text(0.50, 0.57, "Sequential Branch\n(4× GATv2Conv)", ha="center",
            va="bottom", fontsize=8, color="#E67E22", fontweight="bold")

    # Virtual node (side of spatial branch)
    box(ax, 0.82, 0.43, 0.14, 0.05,
        "Virtual Node", "vnode_proj + gate",
        fc="#EFFFEF", ec="#27AE60")
    ax.annotate("", xy=(0.755, 0.43), xytext=(0.75, 0.43),
                arrowprops=dict(arrowstyle="<->", color="#27AE60", lw=1.1,
                                mutation_scale=10), zorder=2)
    ax.plot([0.755, 0.32], [0.43, 0.43], color="#27AE60",
            lw=0.9, linestyle="--", zorder=1)

    # Concatenate outputs
    y_concat = 0.18
    ax.plot([0.20, 0.20], [0.26 - bh/2, y_concat + 0.03], color="#444", lw=1.3)
    ax.plot([0.50, 0.50], [0.33 - bh/2, y_concat + 0.03], color="#444", lw=1.3)
    ax.plot([0.20, 0.50], [y_concat + 0.03, y_concat + 0.03], color="#444", lw=1.3)
    ax.annotate("", xy=(0.35, y_concat + 0.005), xytext=(0.35, y_concat + 0.03),
                arrowprops=dict(arrowstyle="-|>", color="#444", lw=1.3,
                                mutation_scale=12), zorder=2)

    box(ax, 0.35, y_concat - 0.01, 0.24, bh,
        "Concatenate + Fusion", "Linear(1536→768) → ReLU → Linear(768→1)",
        fc="#DDE8FF", ec=NAVY, bold=True)

    # Output
    arrow(ax, 0.35, y_concat - 0.01 - bh/2, 0.35, 0.055)
    box(ax, 0.35, 0.04, 0.22, 0.055,
        "Sigmoid → Per-residue score ∈ [0, 1]", None,
        fc="#1F497D", ec=NAVY, fontsize=9, bold=True)
    ax.texts[-1].set_color("white")

    # Legend
    legend_elements = [
        mpatches.Patch(fc="#E3EAF8", ec=BLUE, label="Spatial GATv2"),
        mpatches.Patch(fc="#FDE8D8", ec="#E67E22", label="Sequential GATv2"),
        mpatches.Patch(fc="#E8F4FD", ec=BLUE, label="Hand-crafted (40-dim)"),
        mpatches.Patch(fc="#FFF3E0", ec="#E67E22", label="ESM2 (480-dim)"),
        mpatches.Patch(fc="#EFFFEF", ec="#27AE60", label="Virtual node"),
    ]
    ax.legend(handles=legend_elements, loc="lower right",
              fontsize=7.5, framealpha=0.9, edgecolor="#cccccc")

    ax.text(0.01, 0.01, "~13.4M parameters  |  input: 520-dim nodes + 6-dim edges",
            fontsize=7.5, color=DGRAY, va="bottom")

    plt.tight_layout(pad=0.3)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=200, bbox_inches="tight",
                facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


# ──────────────────────────────────────────────────────────────────────────────
# Figure 2 — Inference Pipeline
# ──────────────────────────────────────────────────────────────────────────────

def make_pipeline_diagram() -> bytes:
    fig, ax = plt.subplots(figsize=(14, 5.5))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    ax.text(0.5, 0.96, "GNN-PCNA Inference Pipeline",
            ha="center", va="top", fontsize=13, fontweight="bold", color=NAVY)

    # Step boxes (y=0.6 main row), (y=0.3 sub-row for parallel ESM2 path)
    steps = [
        # (cx, cy, w, h, label, sublabel, fc, ec)
        (0.07, 0.60, 0.10, 0.16, "PDB File", ".pdb / RCSB fetch", "#F5F5F5", "#888"),
        (0.21, 0.60, 0.12, 0.16, "Structure\nParser",  "BioPython\nAtom extraction", "#E8F4FD", BLUE),
        (0.36, 0.60, 0.12, 0.16, "Graph\nBuilder",    "k-NN (k=10, r=10Å)\n6-dim edges", "#E8F4FD", BLUE),
        (0.52, 0.60, 0.12, 0.16, "Hand-crafted\nFeatures", "40-dim nodes\n(SASA, SS, φ/ψ…)", "#E3EAF8", BLUE),
    ]
    esm_steps = [
        (0.36, 0.28, 0.12, 0.16, "Sequence\nExtract",  "FASTA from\nstructure", "#FFF3E0", "#E67E22"),
        (0.52, 0.28, 0.12, 0.16, "ESM2\n35M Model",   "facebook/esm2\n480-dim embed.", "#FDE8D8", "#E67E22"),
    ]
    concat_step = (0.67, 0.46, 0.12, 0.16, "Concatenate", "520-dim input\nper residue", "#DDE8FF", NAVY)
    model_step  = (0.81, 0.46, 0.11, 0.16, "PocketGNNXL", "13.4M params\nforward pass", "#1F497D", NAVY)
    post_steps  = [
        (0.81, 0.20, 0.11, 0.16, "DBSCAN\nCluster",   "eps=6Å\nmin_pts=3", "#E8F4FD", BLUE),
        (0.93, 0.20, 0.09, 0.16, "Output", "Pocket\ncandidates", "#F0FFF0", "#27AE60"),
    ]

    def draw_box(s):
        cx, cy, w, h = s[0], s[1], s[2], s[3]
        lbl, sub, fc, ec = s[4], s[5], s[6], s[7]
        rect = FancyBboxPatch((cx-w/2, cy-h/2), w, h,
                              boxstyle="round,pad=0.015", lw=1.4,
                              edgecolor=ec, facecolor=fc, zorder=3)
        ax.add_patch(rect)
        lines = lbl.split("\n")
        if sub:
            ax.text(cx, cy + 0.025, lbl, ha="center", va="center",
                    fontsize=8.5, fontweight="bold",
                    color="white" if fc == "#1F497D" else "#1a1a1a",
                    zorder=4, multialignment="center")
            ax.text(cx, cy - 0.038, sub, ha="center", va="center",
                    fontsize=7, color="#aaa" if fc == "#1F497D" else DGRAY,
                    zorder=4, multialignment="center", style="italic")
        else:
            ax.text(cx, cy, lbl, ha="center", va="center",
                    fontsize=8.5, fontweight="bold", color="#1a1a1a",
                    zorder=4, multialignment="center")

    for s in steps:
        draw_box(s)
    for s in esm_steps:
        draw_box(s)
    draw_box(concat_step)
    draw_box(model_step)
    for s in post_steps:
        draw_box(s)

    # Main row arrows
    for i in range(len(steps) - 1):
        x1 = steps[i][0] + steps[i][2]/2
        x2 = steps[i+1][0] - steps[i+1][2]/2
        y  = steps[i][1]
        arrow(ax, x1, y, x2, y)

    # 40-dim → concatenate
    arrow(ax, steps[3][0] + steps[3][2]/2, 0.60, concat_step[0] - concat_step[2]/2, concat_step[1])

    # ESM2 path: from graph builder down
    ax.plot([0.36, 0.36], [0.60 - 0.08, 0.28 + 0.08],
            color="#E67E22", lw=1.2, linestyle="--", zorder=2)
    ax.annotate("", xy=(0.36, 0.36), xytext=(0.36, 0.52),
                arrowprops=dict(arrowstyle="-|>", color="#E67E22", lw=1.2,
                                mutation_scale=11), zorder=2)
    # ESM2 row arrows
    arrow(ax, esm_steps[0][0] + esm_steps[0][2]/2, 0.28,
               esm_steps[1][0] - esm_steps[1][2]/2, 0.28, col="#E67E22")
    # ESM2 → concatenate
    arrow(ax, esm_steps[1][0] + esm_steps[1][2]/2, 0.28,
               concat_step[0] - concat_step[2]/2, concat_step[1], col="#E67E22")

    # Cache bypass note
    ax.text(0.595, 0.24, "cached\ndata/esm_features/", ha="center", va="center",
            fontsize=6.5, color="#E67E22", style="italic",
            bbox=dict(fc="white", ec="#E67E22", boxstyle="round,pad=0.15", lw=0.7))

    # Concat → model
    arrow(ax, concat_step[0] + concat_step[2]/2, concat_step[1],
               model_step[0]  - model_step[2]/2,  model_step[1])

    # Model → DBSCAN (down)
    ax.plot([model_step[0], model_step[0]],
            [model_step[1] - model_step[3]/2, post_steps[0][1] + post_steps[0][3]/2],
            color="#444", lw=1.3, zorder=2)
    ax.annotate("", xy=(post_steps[0][0], post_steps[0][1] + post_steps[0][3]/2),
                xytext=(model_step[0], post_steps[0][1] + post_steps[0][3]/2),
                arrowprops=dict(arrowstyle="-|>", color="#444", lw=1.3,
                                mutation_scale=12), zorder=2)

    # DBSCAN → output
    arrow(ax, post_steps[0][0] + post_steps[0][2]/2, post_steps[0][1],
               post_steps[1][0] - post_steps[1][2]/2, post_steps[1][1])

    # Threshold annotation
    ax.text(model_step[0] + 0.003, model_step[1] - model_step[3]/2 - 0.04,
            "threshold\n> 0.40", ha="center", va="top",
            fontsize=6.5, color=DGRAY, style="italic")

    # Labels
    ax.text(0.35, 0.85, "Graph Construction", ha="center", fontsize=8,
            color=BLUE, fontweight="bold", style="italic")
    ax.text(0.44, 0.13, "ESM2 Language Model Path (pre-computed or on-the-fly)",
            ha="center", fontsize=7.5, color="#E67E22", style="italic")
    ax.text(0.81, 0.82, "Inference", ha="center", fontsize=8,
            color=NAVY, fontweight="bold", style="italic")
    ax.text(0.87, 0.07, "Post-processing", ha="center", fontsize=8,
            color="#27AE60", fontweight="bold", style="italic")

    # Legend
    legend_elements = [
        mpatches.Patch(fc="#E8F4FD", ec=BLUE, label="Graph / structural"),
        mpatches.Patch(fc="#FDE8D8", ec="#E67E22", label="ESM2 language model"),
        mpatches.Patch(fc="#1F497D", ec=NAVY, label="PocketGNNXL inference"),
        mpatches.Patch(fc="#F0FFF0", ec="#27AE60", label="Output"),
    ]
    ax.legend(handles=legend_elements, loc="upper right",
              fontsize=7.5, framealpha=0.9, edgecolor="#cccccc")

    plt.tight_layout(pad=0.3)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


# ──────────────────────────────────────────────────────────────────────────────
# Word document builder
# ──────────────────────────────────────────────────────────────────────────────

def add_figure(doc, img_bytes_or_path, caption_number: str,
               caption_title: str, caption_body: str,
               width_cm: float = 15.5, landscape: bool = False):
    """Insert a figure with a caption into the document."""
    # Figure image
    if isinstance(img_bytes_or_path, (str, Path)):
        img_stream = str(img_bytes_or_path)
    else:
        img_stream = io.BytesIO(img_bytes_or_path)

    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(img_stream, width=Cm(width_cm))

    # Caption paragraph
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(3)
    p_cap.paragraph_format.space_after = Pt(14)

    # "Figure N — Title" in bold navy
    run_label = p_cap.add_run(f"Figure {caption_number} — {caption_title}")
    run_label.bold = True
    run_label.font.size = Pt(8.5)
    run_label.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    run_label.font.name = "Arial"

    # Body text
    if caption_body:
        run_body = p_cap.add_run(f"  {caption_body}")
        run_body.font.size = Pt(8.5)
        run_body.font.name = "Times New Roman"
        run_body.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    # Horizontal rule after caption
    p_rule = doc.add_paragraph()
    p_rule.paragraph_format.space_before = Pt(0)
    p_rule.paragraph_format.space_after = Pt(4)
    pPr = p_rule._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_page_margins(doc, top=2.0, bottom=2.0, left=2.0, right=2.0):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)


def add_section_header(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    # Blue bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    pBdr.append(bottom)
    pPr.append(pBdr)


def build_doc():
    doc = Document()
    set_page_margins(doc, top=2.2, bottom=2.2, left=2.2, right=2.2)

    # ── Cover block ─────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("GNN-PCNA  |  Figures & Architecture Diagrams")
    r.bold = True
    r.font.size = Pt(15)
    r.font.name = "Arial"
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(
        "Source document for journal figure assembly. "
        "All figures are publication-quality (200–300 DPI). "
        "Figures 1–2 are generated architecture diagrams; "
        "Figures 3–7 are computed from inference results."
    )
    r2.font.size = Pt(9)
    r2.font.name = "Times New Roman"
    r2.font.color.rgb = RGBColor(0x50, 0x50, 0x50)
    doc.add_paragraph()  # spacer

    # ── Table of contents ───────────────────────────────────────────────────
    add_section_header(doc, "Contents")
    toc = [
        ("1", "PocketGNNXL Model Architecture",
         "Dual-branch GATv2 with virtual node, 13.4M params, 520-dim input"),
        ("2", "GNN-PCNA Inference Pipeline",
         "End-to-end flow from PDB file to pocket candidates"),
        ("3", "Comprehensive Results Summary (V1 vs V3)",
         "5-panel overview: AUROC, score distributions, residue profiles, scatter"),
        ("4", "V3 Full Evaluation — 90 PCNA & CryptoSite Structures",
         "Peak pocket score landscape, AUROC distribution (mean 0.940)"),
        ("5", "PCNA Deep Dive — Holo vs Apo vs Benchmark",
         "8GLA/1W60 per-residue profiles, structure comparison, CryptoSite top-20"),
        ("6", "Per-Structure Analysis — 59 PCNA Structures",
         "Ranked pocket scores, AUROC labeling correction, GT recovery"),
        ("7", "Prediction Validation — Confidence & GT Recovery",
         "Raw vs confidence-adjusted scores, GT pocket hits in top-10 predictions"),
    ]
    for num, title, desc in toc:
        p_toc = doc.add_paragraph(style="List Bullet")
        p_toc.paragraph_format.space_before = Pt(1)
        p_toc.paragraph_format.space_after  = Pt(1)
        r_num = p_toc.add_run(f"Figure {num} — ")
        r_num.bold = True
        r_num.font.name = "Arial"
        r_num.font.size = Pt(8.5)
        r_title = p_toc.add_run(title + ": ")
        r_title.bold = True
        r_title.font.name = "Times New Roman"
        r_title.font.size = Pt(8.5)
        r_desc = p_toc.add_run(desc)
        r_desc.font.name = "Times New Roman"
        r_desc.font.size = Pt(8.5)
        r_desc.font.color.rgb = RGBColor(0x50, 0x50, 0x50)

    doc.add_page_break()

    # ── Figure 1: Architecture ───────────────────────────────────────────────
    add_section_header(doc, "Architecture Diagrams")
    add_figure(
        doc, FIGS / "arch_model.png",
        caption_number="1",
        caption_title="PocketGNNXL Model Architecture",
        caption_body=(
            "Dual-branch graph attention network. Hand-crafted structural features "
            "(40-dim: amino acid identity, SASA, secondary structure, B-factor, "
            "pseudo-dihedrals, residue density, charge, hydrophobicity) are "
            "concatenated with ESM2 protein language model embeddings (480-dim) "
            "to form a 520-dimensional node representation. "
            "A shared pre-encoder (Linear 520→256→512→768 + LayerNorm) projects "
            "nodes into a common space before branching into a 5-layer spatial "
            "GATv2Conv branch and a 4-layer sequential GATv2Conv branch. "
            "A virtual node aggregates global context. Branch outputs are "
            "concatenated and projected via MLP to a per-residue cryptic pocket "
            "score in [0, 1]. Total: ~13.4M parameters."
        ),
        width_cm=15.5,
    )

    # ── Figure 2: Pipeline ───────────────────────────────────────────────────
    add_figure(
        doc, FIGS / "arch_pipeline.png",
        caption_number="2",
        caption_title="GNN-PCNA Inference Pipeline",
        caption_body=(
            "End-to-end inference pipeline from raw PDB file to ranked pocket "
            "candidates. Structure parsing extracts Cα coordinates and residue "
            "properties via BioPython. The graph builder constructs a k-nearest "
            "neighbor graph (k=10, cutoff 10 Å) with 6-dimensional edge features "
            "(distance, inverse distance, sequence separation, same-chain flag, "
            "backbone flag, cross-chain flag). ESM2 embeddings are pre-cached per "
            "structure (data/esm_features/) or generated on-the-fly. "
            "Post-inference, residues scoring above 0.40 are clustered with DBSCAN "
            "(eps=6 Å, min_samples=3) on Cα coordinates to yield discrete pocket "
            "candidates ranked by cluster mean score."
        ),
        width_cm=15.5,
    )
    doc.add_page_break()

    # ── Figure 3: Comprehensive data visualization ───────────────────────────
    add_section_header(doc, "Results Figures")
    add_figure(
        doc,
        FIGS / "data_visualization.png",
        caption_number="3",
        caption_title="Comprehensive Results Summary — PocketGNN V1 vs PocketGNNXL V3",
        caption_body=(
            "(A) V1 vs V3 AUROC on the 7 drug-like ligand structures; "
            "V3 mean AUROC 0.9067 vs V1 mean 0.6927. "
            "(B) V3 maximum cluster score distribution stratified by AOH1996 "
            "ground-truth pocket recovery tier (24/24, 20-23/24, <20/24). "
            "(C) Per-residue score profile for 8GLA chain A (V1 orange, V3 blue) "
            "with AOH1996 ground-truth residues marked; V3 recovers all 24/24. "
            "(D) AOH overlap vs top cluster score for all V3 structures. "
            "(E) Scatter of V1 vs V3 top cluster scores across all 59 PCNA "
            "structures, colored by AOH1996 ground-truth overlap count."
        ),
        width_cm=15.5,
    )

    # ── Figure 4: Full evaluation landscape ─────────────────────────────────
    add_figure(
        doc,
        DATA / "fig1_score_landscape.png",
        caption_number="4",
        caption_title="V3 Full Evaluation — 90 PCNA & CryptoSite Structures",
        caption_body=(
            "Top-left: peak pocket score per structure across 90 PCNA (red) and "
            "CryptoSite benchmark (blue) structures; threshold 0.4 shown. "
            "Top-right: score distribution by dataset; PCNA structures cluster at "
            "higher scores than CryptoSite background. "
            "Bottom-left: AUROC distribution across 53 labeled CryptoSite structures; "
            "mean AUROC = 0.940. "
            "Bottom-right: pocket coverage (% residues above threshold) vs "
            "structure size in residues; PCNA structures labeled."
        ),
        width_cm=15.5,
    )
    doc.add_page_break()

    # ── Figure 5: PCNA deep dive ─────────────────────────────────────────────
    add_figure(
        doc,
        DATA / "fig2_pcna_deepdive.png",
        caption_number="5",
        caption_title="PCNA Deep Dive — Holo (8GLA) vs Apo (1W60) vs Benchmark",
        caption_body=(
            "Top row: per-residue score profiles for all chains of 8GLA (holo, "
            "AOH1996-bound) and 1W60 (apo, no ligand); model correctly suppresses "
            "scores in the apo structure while maintaining high scores at pocket "
            "residues in the holo. "
            "Middle: PCNA structure comparison — mean score, max score, and % "
            "residues above 0.4 for holo, apo, and apo-like controls. "
            "Bottom-left: AUROC for top-20 CryptoSite structures; mean 0.940. "
            "Bottom-right: predicted pocket count distribution (threshold 0.4) "
            "across all evaluated PCNA structures."
        ),
        width_cm=15.5,
    )
    doc.add_page_break()

    # ── Figure 6: Per-structure analysis ─────────────────────────────────────
    add_figure(
        doc,
        RESULTS / "per_structure" / "full_analysis.png",
        caption_number="6",
        caption_title="Per-Structure Analysis — 59 PCNA Structures",
        caption_body=(
            "Top: all 59 PCNA structures ranked by top cluster pocket score; "
            "holo (red), apo (blue), and other (gray) structures color-coded; "
            "AOH1996 ground-truth overlap count overlaid as scatter. "
            "Middle-left: AUROC before vs after labeling correction (filtering "
            "AGS/ADP cofactor-adjacent residues from ground truth). "
            "Middle-right: top cluster mean score vs AOH1996 ground-truth "
            "recovery count for all 59 structures. "
            "Bottom-left: score profiles for 8GLA (holo) vs 9B8T (Pol ε-PCNA "
            "interface novel site candidate). "
            "Bottom-right: AUROC distribution for drug-like ligand structures only."
        ),
        width_cm=15.5,
    )
    doc.add_page_break()

    # ── Figure 7: Prediction validation ──────────────────────────────────────
    add_figure(
        doc,
        RESULTS / "validation" / "confidence.png",
        caption_number="7",
        caption_title="Prediction Validation — Confidence Adjustment & GT Recovery",
        caption_body=(
            "Left: raw top cluster mean score vs confidence-adjusted score for all "
            "PCNA structures; color = confidence. Known holo structures (8GLA, 8GL9) "
            "labeled; calibration is tight for high-confidence predictions. "
            "Center: stacked bar of confidence components per structure — "
            "uncertainty (1-std, blue) and symmetry correction (green); structures "
            "ordered by decreasing adjusted score. "
            "Right: AOH1996 ground-truth hits in the top-10 model predictions vs "
            "confidence-adjusted score; 8GLA and 3VKX recover the most GT residues "
            "within the top-10 list, consistent with highest AUROC values."
        ),
        width_cm=15.5,
    )

    doc.save(str(OUT))
    print(f"Saved -> {OUT}")


if __name__ == "__main__":
    build_doc()
